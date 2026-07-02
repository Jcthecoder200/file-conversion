"""
File Converter
==============
A small desktop tool (Windows) that lets you pick a file, choose a sensible
target format from a filtered drop-down list, and convert it. Also includes
a 4-corner "slice" tool for images and batch folder conversion.

Layout (top to bottom):
  Row 1: [Select File] button  ->  shows chosen file name + type
  Row 2: [source type]  --ASCII arrow-->  [target type dropdown]
  Row 3: output name entry, [Browse] output folder, [Batch Convert] (left),
         [Slice Image] (middle), [Convert] (right)

--------------------------------------------------------------------------
INSTALL (run once, in a Command Prompt / PowerShell):

    pip install -r requirements.txt

requirements.txt contents:
    Pillow
    PyMuPDF
    python-docx
    pdf2docx
    docx2pdf
    fpdf2

Notes:
  - docx2pdf needs Microsoft Word installed on Windows (it drives Word
    in the background). If Word isn't installed, .docx -> .pdf will fail
    with a clear error message instead of crashing the app.
  - Every conversion function imports its own library "lazily", so the
    app still opens and runs even if you're missing a package - you'll
    just get a friendly error only when you try that specific conversion.
--------------------------------------------------------------------------
"""

import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# --------------------------------------------------------------------------
# Colours / style (design protocol: mild background, greyish buttons,
# darker grey when pressed, black text, no bright/shape colours)
# --------------------------------------------------------------------------
BG_COLOR = "#f0f0ee"          # mild background
PANEL_COLOR = "#f0f0ee"
BUTTON_BG = "#d9d9d9"         # greyish button
BUTTON_ACTIVE_BG = "#adadad"  # darker grey when pressed
BUTTON_FG = "black"
TEXT_COLOR = "#222222"
ENTRY_BG = "#ffffff"
BORDER_COLOR = "#bfbfbf"
FONT_NORMAL = ("Segoe UI", 10)
FONT_BOLD = ("Segoe UI", 10, "bold")
FONT_MONO = ("Consolas", 14, "bold")   # for ASCII arrow

IMAGE_EXTS = [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".webp"]

# --------------------------------------------------------------------------
# CASE STUDY / CONVERSION MAP
# Only lists combinations that make practical sense.
# e.g. .png -> .docx is intentionally excluded (would need OCR, not a
# straightforward format conversion), same for .png -> .pdf which IS
# allowed since an image can simply become a page of a pdf.
# --------------------------------------------------------------------------
CONVERSION_MAP = {
    ".pdf":  [".docx", ".png", ".jpg", ".txt"],
    ".docx": [".pdf", ".txt"],
    ".txt":  [".docx", ".pdf"],
    ".png":  [".jpg", ".bmp", ".gif", ".tiff", ".webp", ".pdf"],
    ".jpg":  [".png", ".bmp", ".gif", ".tiff", ".webp", ".pdf"],
    ".jpeg": [".png", ".bmp", ".gif", ".tiff", ".webp", ".pdf"],
    ".bmp":  [".png", ".jpg", ".gif", ".tiff", ".webp", ".pdf"],
    ".gif":  [".png", ".jpg", ".bmp", ".tiff", ".webp", ".pdf"],
    ".tiff": [".png", ".jpg", ".bmp", ".gif", ".webp", ".pdf"],
    ".webp": [".png", ".jpg", ".bmp", ".gif", ".tiff", ".pdf"],
}

FRIENDLY_NAME = {
    ".pdf": "PDF Document", ".docx": "Word Document", ".txt": "Text File",
    ".png": "PNG Image", ".jpg": "JPEG Image", ".jpeg": "JPEG Image",
    ".bmp": "Bitmap Image", ".gif": "GIF Image", ".tiff": "TIFF Image",
    ".webp": "WEBP Image",
}


def get_ext(path):
    return os.path.splitext(path)[1].lower()


def friendly(ext):
    return FRIENDLY_NAME.get(ext, ext.upper().lstrip("."))


# ==========================================================================
# CONVERSION FUNCTIONS  (each imports its own dependency lazily)
# ==========================================================================

def _pillow_save(src_path, dst_path, target_ext):
    from PIL import Image
    im = Image.open(src_path)
    fmt_map = {".jpg": "JPEG", ".jpeg": "JPEG", ".png": "PNG", ".bmp": "BMP",
               ".gif": "GIF", ".tiff": "TIFF", ".webp": "WEBP", ".pdf": "PDF"}
    fmt = fmt_map[target_ext]
    if fmt in ("JPEG", "PDF") and im.mode in ("RGBA", "P"):
        im = im.convert("RGB")
    im.save(dst_path, fmt)


def convert_image_to_image_or_pdf(src_path, dst_path, target_ext=None):
    _pillow_save(src_path, dst_path, target_ext)


def convert_pdf_to_image(src_path, dst_path, target_ext=None):
    """Renders every page; if the pdf has 1 page, dst_path is used as-is,
    otherwise pages are numbered next to the requested file name."""
    import fitz  # PyMuPDF
    doc = fitz.open(src_path)
    base, ext = os.path.splitext(dst_path)
    if len(doc) == 1:
        pix = doc[0].get_pixmap(dpi=200)
        pix.save(dst_path)
    else:
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=200)
            pix.save(f"{base}_page{i}{ext}")
    doc.close()


def convert_pdf_to_docx(src_path, dst_path, target_ext=None):
    from pdf2docx import Converter
    cv = Converter(src_path)
    cv.convert(dst_path)
    cv.close()


def convert_pdf_to_txt(src_path, dst_path, target_ext=None):
    import fitz
    doc = fitz.open(src_path)
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    with open(dst_path, "w", encoding="utf-8") as f:
        f.write(text)


def convert_docx_to_pdf(src_path, dst_path, target_ext=None):
    try:
        from docx2pdf import convert
    except ImportError as e:
        raise RuntimeError("docx2pdf is not installed. Run: pip install docx2pdf") from e
    try:
        convert(src_path, dst_path)
    except Exception as e:
        raise RuntimeError(
            "Could not convert .docx to .pdf. This feature needs Microsoft "
            "Word installed on this Windows PC. Details: " + str(e)
        ) from e


def convert_docx_to_txt(src_path, dst_path, target_ext=None):
    import docx
    d = docx.Document(src_path)
    text = "\n".join(p.text for p in d.paragraphs)
    with open(dst_path, "w", encoding="utf-8") as f:
        f.write(text)


def convert_txt_to_docx(src_path, dst_path, target_ext=None):
    import docx
    d = docx.Document()
    with open(src_path, "r", encoding="utf-8", errors="replace") as f:
        for line in f:
            d.add_paragraph(line.rstrip("\n"))
    d.save(dst_path)


def convert_txt_to_pdf(src_path, dst_path, target_ext=None):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    with open(src_path, "r", encoding="utf-8", errors="replace") as f:
        for line in f:
            pdf.multi_cell(0, 8, line.rstrip("\n"))
    pdf.output(dst_path)


# dispatch table keyed by (source_ext, target_ext)
def build_dispatch():
    table = {}
    for img_ext in IMAGE_EXTS:
        for tgt in CONVERSION_MAP.get(img_ext, []):
            if tgt in IMAGE_EXTS or tgt == ".pdf":
                table[(img_ext, tgt)] = convert_image_to_image_or_pdf
    table[(".pdf", ".png")] = convert_pdf_to_image
    table[(".pdf", ".jpg")] = convert_pdf_to_image
    table[(".pdf", ".docx")] = convert_pdf_to_docx
    table[(".pdf", ".txt")] = convert_pdf_to_txt
    table[(".docx", ".pdf")] = convert_docx_to_pdf
    table[(".docx", ".txt")] = convert_docx_to_txt
    table[(".txt", ".docx")] = convert_txt_to_docx
    table[(".txt", ".pdf")] = convert_txt_to_pdf
    return table


DISPATCH = build_dispatch()


# ==========================================================================
# GUI helpers
# ==========================================================================

def make_button(parent, text, command, width=16):
    return tk.Button(
        parent, text=text, command=command, width=width,
        bg=BUTTON_BG, fg=BUTTON_FG, activebackground=BUTTON_ACTIVE_BG,
        activeforeground=BUTTON_FG, relief="raised", bd=1,
        font=FONT_NORMAL, cursor="hand2", highlightthickness=0,
    )


# ==========================================================================
# Image slicer (4 draggable corner dots -> crop/warp -> save/convert)
# ==========================================================================

class ImageSlicer(tk.Toplevel):
    DOT_R = 7
    MAX_DIM = 560

    def __init__(self, master, src_path, target_ext, on_done):
        super().__init__(master)
        self.title("Slice Image")
        self.configure(bg=BG_COLOR)
        self.resizable(False, False)
        self.src_path = src_path
        self.target_ext = target_ext
        self.on_done = on_done

        from PIL import Image, ImageTk
        self.Image = Image
        self.ImageTk = ImageTk

        self.pil_img = Image.open(src_path)
        self.pil_img = self.pil_img.convert("RGB") if self.pil_img.mode not in ("RGB", "RGBA") else self.pil_img
        w, h = self.pil_img.size
        scale = min(self.MAX_DIM / w, self.MAX_DIM / h, 1.0)
        self.scale = scale
        disp_w, disp_h = int(w * scale), int(h * scale)
        self.disp_size = (disp_w, disp_h)

        self.tk_img = ImageTk.PhotoImage(self.pil_img.resize((disp_w, disp_h)))

        tk.Label(self, text="Drag the 4 dots to the corners you want to keep, then confirm.",
                 bg=BG_COLOR, fg=TEXT_COLOR, font=FONT_NORMAL).pack(pady=(10, 4))

        self.canvas = tk.Canvas(self, width=disp_w, height=disp_h,
                                 bg="#ffffff", highlightthickness=1,
                                 highlightbackground=BORDER_COLOR)
        self.canvas.pack(padx=10, pady=6)
        self.canvas.create_image(0, 0, anchor="nw", image=self.tk_img)

        margin = 20
        # order required by PIL QUAD transform: upper-left, lower-left,
        # lower-right, upper-right
        self.points = [
            [margin, margin],
            [margin, disp_h - margin],
            [disp_w - margin, disp_h - margin],
            [disp_w - margin, margin],
        ]
        self.dot_ids = []
        self.line_id = None
        self._draw_shape()

        for idx in range(4):
            self.canvas.tag_bind(self.dot_ids[idx], "<B1-Motion>",
                                  lambda e, i=idx: self._drag(e, i))

        btn_row = tk.Frame(self, bg=BG_COLOR)
        btn_row.pack(pady=10)
        make_button(btn_row, "Cancel", self.destroy).pack(side="left", padx=6)
        make_button(btn_row, "Confirm Slice && Convert", self._confirm, width=22).pack(side="left", padx=6)

    def _draw_shape(self):
        if self.line_id:
            self.canvas.delete(self.line_id)
        pts = self.points + [self.points[0]]
        flat = [c for p in pts for c in p]
        self.line_id = self.canvas.create_line(*flat, fill="#555555", width=2)
        self.canvas.tag_lower(self.line_id)
        if not self.dot_ids:
            for (x, y) in self.points:
                dot = self.canvas.create_oval(
                    x - self.DOT_R, y - self.DOT_R, x + self.DOT_R, y + self.DOT_R,
                    fill="#9c9c9c", outline="#4d4d4d", width=2
                )
                self.dot_ids.append(dot)
        else:
            for dot, (x, y) in zip(self.dot_ids, self.points):
                self.canvas.coords(dot, x - self.DOT_R, y - self.DOT_R, x + self.DOT_R, y + self.DOT_R)

    def _drag(self, event, idx):
        w, h = self.disp_size
        x = min(max(event.x, 0), w)
        y = min(max(event.y, 0), h)
        self.points[idx] = [x, y]
        self._draw_shape()

    def _confirm(self):
        # map displayed coords back to original image coords
        src_pts = [(x / self.scale, y / self.scale) for (x, y) in self.points]
        ul, ll, lr, ur = src_pts

        def dist(a, b):
            return ((a[0] - b[0]) ** 2 + (a[1] - b[1]) ** 2) ** 0.5

        out_w = int(max(dist(ul, ur), dist(ll, lr)))
        out_h = int(max(dist(ul, ll), dist(ur, lr)))
        out_w = max(out_w, 10)
        out_h = max(out_h, 10)

        quad_data = [ul[0], ul[1], ll[0], ll[1], lr[0], lr[1], ur[0], ur[1]]
        try:
            sliced = self.pil_img.transform(
                (out_w, out_h), self.Image.QUAD, quad_data, resample=self.Image.BICUBIC
            )
        except Exception as e:
            messagebox.showerror("Slice failed", str(e))
            return
        self.destroy()
        self.on_done(sliced)


# ==========================================================================
# Batch Converter Dialog
# ==========================================================================

class BatchConverterDialog(tk.Toplevel):
    def __init__(self, master):
        super().__init__(master)
        self.master = master
        self.title("Batch Convert Folder")
        self.configure(bg=BG_COLOR)
        self.resizable(False, False)
        self.geometry("500x300")

        self.src_folder = tk.StringVar()
        self.out_folder = tk.StringVar()
        self.target_ext = tk.StringVar()

        # Build UI
        main = tk.Frame(self, bg=BG_COLOR)
        main.pack(fill="both", expand=True, padx=16, pady=16)

        # Source folder
        tk.Label(main, text="Source Folder:", bg=BG_COLOR, fg=TEXT_COLOR, font=FONT_NORMAL).grid(row=0, column=0, sticky="w", pady=4)
        src_entry = tk.Entry(main, textvariable=self.src_folder, width=40, bg=ENTRY_BG, fg=TEXT_COLOR, relief="solid", bd=1, font=FONT_NORMAL)
        src_entry.grid(row=0, column=1, padx=6, sticky="ew")
        make_button(main, "Browse...", self._browse_src, width=10).grid(row=0, column=2, padx=4)

        # Output folder
        tk.Label(main, text="Output Folder:", bg=BG_COLOR, fg=TEXT_COLOR, font=FONT_NORMAL).grid(row=1, column=0, sticky="w", pady=4)
        out_entry = tk.Entry(main, textvariable=self.out_folder, width=40, bg=ENTRY_BG, fg=TEXT_COLOR, relief="solid", bd=1, font=FONT_NORMAL)
        out_entry.grid(row=1, column=1, padx=6, sticky="ew")
        make_button(main, "Browse...", self._browse_out, width=10).grid(row=1, column=2, padx=4)

        # Target type
        tk.Label(main, text="Target Type:", bg=BG_COLOR, fg=TEXT_COLOR, font=FONT_NORMAL).grid(row=2, column=0, sticky="w", pady=10)
        # All possible target extensions
        all_targets = sorted({ext for targets in CONVERSION_MAP.values() for ext in targets})
        self.target_combo = ttk.Combobox(main, textvariable=self.target_ext, values=all_targets, state="readonly", width=20, font=FONT_NORMAL)
        self.target_combo.grid(row=2, column=1, sticky="w", padx=6)
        if all_targets:
            self.target_combo.set(all_targets[0])

        # Status label
        self.status = tk.Label(main, text="", bg=BG_COLOR, fg=TEXT_COLOR, font=FONT_NORMAL, anchor="w")
        self.status.grid(row=3, column=0, columnspan=3, sticky="w", pady=10)

        # Start button
        self.start_btn = make_button(main, "Start Batch", self._start_batch, width=16)
        self.start_btn.grid(row=4, column=0, columnspan=3, pady=10)

        # Configure grid weights
        main.columnconfigure(1, weight=1)

    def _browse_src(self):
        folder = filedialog.askdirectory(title="Select source folder")
        if folder:
            self.src_folder.set(folder)
            # Auto-fill output folder if not set
            if not self.out_folder.get():
                self.out_folder.set(folder + "_converted")

    def _browse_out(self):
        folder = filedialog.askdirectory(title="Select output folder")
        if folder:
            self.out_folder.set(folder)

    def _start_batch(self):
        src = self.src_folder.get().strip()
        out = self.out_folder.get().strip()
        target = self.target_ext.get()
        if not src or not os.path.isdir(src):
            messagebox.showerror("Error", "Please select a valid source folder.")
            return
        if not out:
            messagebox.showerror("Error", "Please select an output folder.")
            return
        if not target:
            messagebox.showerror("Error", "Please select a target file type.")
            return

        # Create output folder if needed
        os.makedirs(out, exist_ok=True)

        self.start_btn.config(state="disabled")
        self.status.config(text="Processing...")

        # Collect files
        files = [f for f in os.listdir(src) if os.path.isfile(os.path.join(src, f))]
        total = 0
        converted = 0
        errors = []

        for filename in files:
            src_path = os.path.join(src, filename)
            src_ext = get_ext(src_path)
            if src_ext not in CONVERSION_MAP:
                continue
            if target not in CONVERSION_MAP[src_ext]:
                continue  # skip if conversion not supported

            total += 1
            # Build output filename: keep basename, change extension
            base = os.path.splitext(filename)[0]
            dst_filename = base + target
            dst_path = os.path.join(out, dst_filename)

            # Avoid overwriting: add _1, _2, ...
            counter = 1
            while os.path.exists(dst_path):
                dst_filename = f"{base}_{counter}{target}"
                dst_path = os.path.join(out, dst_filename)
                counter += 1

            try:
                func = DISPATCH.get((src_ext, target))
                if func is None:
                    errors.append(f"❌ {filename}: no converter registered")
                    continue
                func(src_path, dst_path, target)
                converted += 1
                self.status.config(text=f"Converted {converted}/{total}...")
                self.update_idletasks()
            except Exception as e:
                errors.append(f"❌ {filename}: {str(e)}")

        # Final summary
        msg = f"Done! Converted {converted} file(s)."
        if errors:
            msg += f"\n\n{len(errors)} error(s):\n" + "\n".join(errors[:10])
            if len(errors) > 10:
                msg += f"\n... and {len(errors)-10} more."
        messagebox.showinfo("Batch Complete", msg)
        self.status.config(text="Ready")
        self.start_btn.config(state="normal")
        self.destroy()


# ==========================================================================
# Main application
# ==========================================================================

class ConverterApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("File Converter")
        self.configure(bg=BG_COLOR)
        self.resizable(False, False)
        self.geometry("620x320")

        self.src_path = tk.StringVar()
        self.src_ext = None
        self.out_folder = tk.StringVar()
        self.out_name = tk.StringVar()

        self._build_row1()
        self._build_row2()
        self._build_row3()
        self._refresh_state()

    # ---- row 1 --------------------------------------------------------
    def _build_row1(self):
        row = tk.Frame(self, bg=BG_COLOR)
        row.pack(fill="x", padx=16, pady=(16, 8))

        make_button(row, "Select File...", self._select_file, width=16).pack(side="left")
        self.file_label = tk.Label(
            row, text="No file selected", bg=BG_COLOR, fg=TEXT_COLOR,
            font=FONT_NORMAL, anchor="w"
        )
        self.file_label.pack(side="left", padx=12, fill="x", expand=True)

    # ---- row 2 --------------------------------------------------------
    def _build_row2(self):
        row = tk.Frame(self, bg=BG_COLOR)
        row.pack(fill="x", padx=16, pady=8)

        self.type_left = tk.Label(
            row, text="(select a file)", bg=BG_COLOR, fg=TEXT_COLOR,
            font=FONT_BOLD, width=16, anchor="w"
        )
        self.type_left.pack(side="left")

        arrow = tk.Label(row, text="-------->", bg=BG_COLOR, fg=TEXT_COLOR, font=FONT_MONO)
        arrow.pack(side="left", padx=10)

        self.target_var = tk.StringVar()
        self.target_combo = ttk.Combobox(
            row, textvariable=self.target_var, state="disabled",
            values=[], width=18, font=FONT_NORMAL
        )
        self.target_combo.pack(side="left")
        self.target_combo.bind("<<ComboboxSelected>>", lambda e: self._refresh_state())

    # ---- row 3 --------------------------------------------------------
    def _build_row3(self):
        row = tk.Frame(self, bg=BG_COLOR)
        row.pack(fill="x", padx=16, pady=8)

        tk.Label(row, text="Save as:", bg=BG_COLOR, fg=TEXT_COLOR, font=FONT_NORMAL).grid(row=0, column=0, sticky="w")
        self.name_entry = tk.Entry(row, textvariable=self.out_name, width=20, bg=ENTRY_BG,
                                    fg=TEXT_COLOR, relief="solid", bd=1, font=FONT_NORMAL)
        self.name_entry.grid(row=0, column=1, padx=(6, 12), sticky="w")

        make_button(row, "Choose Folder...", self._choose_folder, width=14).grid(row=0, column=2, sticky="w")
        self.folder_label = tk.Label(row, text="(no folder chosen)", bg=BG_COLOR, fg=TEXT_COLOR,
                                      font=FONT_NORMAL, anchor="w")
        self.folder_label.grid(row=1, column=0, columnspan=3, sticky="w", pady=(6, 0))

        # Buttons row
        btn_row = tk.Frame(self, bg=BG_COLOR)
        btn_row.pack(fill="x", padx=16, pady=(18, 10))

        self.batch_btn = make_button(btn_row, "Batch Convert...", self._open_batch_dialog, width=16)
        self.batch_btn.pack(side="left")

        self.slice_btn = make_button(btn_row, "Slice Image...", self._open_slicer, width=16)
        self.slice_btn.pack(side="left", padx=8)

        self.convert_btn = make_button(btn_row, "Convert", self._convert, width=16)
        self.convert_btn.pack(side="right")

    # ---- behaviour ------------------------------------------------------
    def _select_file(self):
        path = filedialog.askopenfilename(
            title="Select a file to convert",
            filetypes=[("Supported files", "*.pdf *.docx *.txt *.png *.jpg *.jpeg *.bmp *.gif *.tiff *.webp"),
                       ("All files", "*.*")]
        )
        if not path:
            return  # user cancelled - keeps previous selection, lets them retry
        ext = get_ext(path)
        if ext not in CONVERSION_MAP:
            messagebox.showwarning(
                "Unsupported file type",
                f"'{ext}' is not supported yet. Supported source types: "
                f"{', '.join(sorted(CONVERSION_MAP.keys()))}"
            )
            return
        self.src_path.set(path)
        self.src_ext = ext
        self.file_label.config(text=f"{os.path.basename(path)}   ({friendly(ext)})")
        self.type_left.config(text=f"Type: {ext}")

        base_name = os.path.splitext(os.path.basename(path))[0]
        self.out_name.set(base_name)
        if not self.out_folder.get():
            self.out_folder.set(os.path.dirname(path))
            self.folder_label.config(text=self.out_folder.get())

        # --- CASE STUDY CHECK: only show sensible targets for this ext ---
        options = CONVERSION_MAP.get(ext, [])
        self.target_combo.config(values=options, state="readonly" if options else "disabled")
        if options:
            self.target_var.set(options[0])
        else:
            self.target_var.set("")

        self._refresh_state()

    def _choose_folder(self):
        folder = filedialog.askdirectory(title="Choose output folder")
        if folder:
            self.out_folder.set(folder)
            self.folder_label.config(text=folder)

    def _refresh_state(self):
        is_image = self.src_ext in IMAGE_EXTS
        self.slice_btn.config(state="normal" if is_image else "disabled")
        can_convert = bool(self.src_ext and self.target_var.get() and self.out_folder.get() and self.out_name.get())
        self.convert_btn.config(state="normal" if can_convert else "disabled")

    def _dest_path(self, target_ext=None):
        target_ext = target_ext or self.target_var.get()
        name = self.out_name.get().strip()
        if not name:
            name = "converted"
        if not name.lower().endswith(target_ext):
            name += target_ext
        return os.path.join(self.out_folder.get(), name)

    def _convert(self):
        target_ext = self.target_var.get()
        if not target_ext:
            messagebox.showwarning("Missing target type", "Please choose a target file type.")
            return
        if target_ext not in CONVERSION_MAP.get(self.src_ext, []):
            messagebox.showerror("Not supported", f"Cannot convert {self.src_ext} to {target_ext}.")
            return
        func = DISPATCH.get((self.src_ext, target_ext))
        if func is None:
            messagebox.showerror("Not implemented", f"No converter registered for {self.src_ext} -> {target_ext}.")
            return
        dst_path = self._dest_path(target_ext)
        try:
            func(self.src_path.get(), dst_path, target_ext)
            messagebox.showinfo("Done", f"Saved to:\n{dst_path}")
        except Exception as e:
            messagebox.showerror("Conversion failed", str(e))

    def _open_slicer(self):
        if self.src_ext not in IMAGE_EXTS:
            return
        target_ext = self.target_var.get() or self.src_ext
        ImageSlicer(self, self.src_path.get(), target_ext, self._on_slice_done)

    def _on_slice_done(self, sliced_pil_image):
        target_ext = self.target_var.get() or self.src_ext
        dst_path = self._dest_path(target_ext)
        try:
            fmt_map = {".jpg": "JPEG", ".jpeg": "JPEG", ".png": "PNG", ".bmp": "BMP",
                       ".gif": "GIF", ".tiff": "TIFF", ".webp": "WEBP", ".pdf": "PDF"}
            fmt = fmt_map.get(target_ext, "PNG")
            img = sliced_pil_image
            if fmt in ("JPEG", "PDF") and img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(dst_path, fmt)
            messagebox.showinfo("Done", f"Sliced image saved to:\n{dst_path}")
        except Exception as e:
            messagebox.showerror("Save failed", str(e))

    def _open_batch_dialog(self):
        BatchConverterDialog(self)


if __name__ == "__main__":
    app = ConverterApp()
    app.mainloop()